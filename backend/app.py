import re
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import cv2
from docx import Document

# PATHS
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\Users\vikas\Downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin'

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# 🔍 ADVANCED DETECTION
def detect_sensitive_data(text):
    text = text.replace("\n", " ")

    patterns = {
        "Aadhaar": r'\b\d{4}\s?\d{4}\s?\d{4}\b',
        "PAN": r'\b[A-Z]{5}[0-9]{4}[A-Z]\b',
        "Phone": r'\b[6-9]\d{9}\b',
        "Email": r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b',
        "Credit Card": r'\b(?:\d{4}[-\s]?){3}\d{4}\b',
        "Passport": r'\b[A-Z][0-9]{7}\b',
        "IFSC": r'\b[A-Z]{4}0[A-Z0-9]{6}\b',
        "Bank Account": r'\b\d{9,18}\b'
    }

    detected = {}

    for key, pattern in patterns.items():
        matches = re.findall(pattern, text)
        if matches:
            detected[key] = list(set(matches))

    return detected


# 🖼 MASK IMAGE
def mask_image(filepath, values):
    img = cv2.imread(filepath)
    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    chunks = []
    for v in values:
        clean = v.replace(" ", "")
        chunks += [clean[i:i+4] for i in range(0, len(clean), 4)]

    for i in range(len(data['text'])):
        word = data['text'][i]
        for chunk in chunks:
            if chunk and chunk in word:
                x = data['left'][i]
                y = data['top'][i]
                w = data['width'][i]
                h = data['height'][i]
                cv2.rectangle(img, (x,y), (x+w,y+h), (0,0,0), -1)

    out = "masked_" + os.path.basename(filepath)
    out_path = os.path.join(UPLOAD_FOLDER, out)
    cv2.imwrite(out_path, img)

    return out


# 📷 PREVIEW
@app.route('/image/<filename>')
def image(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)


# ⬇ DOWNLOAD
@app.route('/download/<filename>')
def download(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)


@app.route('/')
def home():
    return "Backend running 🚀"


# 🚀 MAIN
@app.route('/upload', methods=['POST'])
def upload():
    try:
        file = request.files['file']
        file_type = request.form.get("type")
        path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(path)

        # 🖼 IMAGE
        if file_type == "image":
            text = pytesseract.image_to_string(Image.open(path))
            detected = detect_sensitive_data(text)

            all_values = []
            for v in detected.values():
                all_values.extend(v)

            masked = mask_image(path, all_values)

            return jsonify({
                "type": "image",
                "detected": detected,
                "preview": f"http://127.0.0.1:5000/image/{masked}",
                "download": f"http://127.0.0.1:5000/download/{masked}"
            })

        # 📄 PDF
        elif file_type == "pdf":
            images = convert_from_path(path, poppler_path=POPPLER_PATH)
            masked_imgs = []
            detected_all = {}

            for i, img in enumerate(images):
                temp = os.path.join(UPLOAD_FOLDER, f"temp_{i}.png")
                img.save(temp)

                text = pytesseract.image_to_string(img)
                detected = detect_sensitive_data(text)

                for k, v in detected.items():
                    detected_all.setdefault(k, []).extend(v)

                all_values = []
                for v in detected.values():
                    all_values.extend(v)

                masked_name = mask_image(temp, all_values)
                masked_imgs.append(Image.open(os.path.join(UPLOAD_FOLDER, masked_name)))

            pdf_path = os.path.join(UPLOAD_FOLDER, "masked.pdf")
            masked_imgs[0].save(pdf_path, save_all=True, append_images=masked_imgs[1:])

            return jsonify({
                "type": "pdf",
                "detected": detected_all,
                "download": "http://127.0.0.1:5000/download/masked.pdf"
            })

        # 📝 TXT
        elif file_type == "txt":
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()

            detected = detect_sensitive_data(text)

            masked_text = text
            for values in detected.values():
                for val in values:
                    masked_text = masked_text.replace(val, "XXXX")

            txt_path = os.path.join(UPLOAD_FOLDER, "masked.txt")
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(masked_text)

            return jsonify({
                "type": "txt",
                "detected": detected,
                "download": "http://127.0.0.1:5000/download/masked.txt"
            })

        # 📄 DOCX
        elif file_type == "docx":
            doc = Document(path)
            text = "\n".join([p.text for p in doc.paragraphs])

            detected = detect_sensitive_data(text)

            masked_text = text
            for values in detected.values():
                for val in values:
                    masked_text = masked_text.replace(val, "XXXX")

            new_doc = Document()
            new_doc.add_paragraph(masked_text)

            docx_path = os.path.join(UPLOAD_FOLDER, "masked.docx")
            new_doc.save(docx_path)

            return jsonify({
                "type": "docx",
                "detected": detected,
                "download": "http://127.0.0.1:5000/download/masked.docx"
            })

        else:
            return jsonify({"error": "Invalid type"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True)